import csv
import io
import os
import re
import zipfile
from flask import send_file
import numpy as np
import torch
import PyPDF2
import docx2txt
import docx
from docx.enum.text import WD_COLOR_INDEX
import zipfile
from transformers import GPT2Tokenizer, GPT2LMHeadModel, BertTokenizer, BertForMaskedLM
from .class_register import register_api

class AbstractLanguageChecker:
    def __init__(self):
        self.device = torch.device(
            "cuda" if torch.cuda.is_available() else "cpu")

    def check_probabilities(self, in_text, topk=40):
        raise NotImplementedError

    def postprocess(self, token):
        raise NotImplementedError

    def top_k_logits(logits, k):
        """
        Filters logits to only the top k choices
        from http://github.com/huggingface/pytorch-pretrained-BERT/blob/master/examples/run_gpt2.py
        """
        if k == 0:
            return logits
        values, _ = torch.topk(logits, k)
        min_values = values[:, -1]
        return torch.where(logits < min_values,
                       torch.ones_like(logits, dtype=logits.dtype) * -1e10,
                       logits)

@register_api(name='gpt-2')
class LM(AbstractLanguageChecker):
    def __init__(self, model_name_or_path="gpt2"):
        super(LM, self).__init__()
        self.enc = GPT2Tokenizer.from_pretrained(model_name_or_path)
        self.model = GPT2LMHeadModel.from_pretrained(model_name_or_path)
        self.model.to(self.device)
        self.model.eval()
        self.start_token = self.enc.bos_token_id
        print("Loaded GPT-2 model!")
    
    def check_probabilities(self, in_text, topk, para):
        word_count = [0,0,0,0]
        # create in-text to change pdf or docx to test
        token_ids = self.enc(in_text, return_tensors='pt').data['input_ids'][0]
        token_ids = torch.concat([self.start_token, token_ids])
        # Forward through the model
        output = self.model(token_ids.to(self.device))
        all_logits = output.logits[:-1].detach().squeeze()
        # construct target and pred
        all_probs = torch.softmax(all_logits, dim=1)

        y = token_ids[1:]
        # Sort the predictions for each timestep
        sorted_preds = torch.argsort(all_probs, dim=1, descending=True).cpu()
        # [(pos, prob), ...]
        real_topk_pos = list(
            [int(np.where(sorted_preds[i] == y[i].item())[0][0])
             for i in range(y.shape[0])])

        bpe_strings = self.enc.convert_ids_to_tokens(token_ids[:])

        bpe_strings = [self.postprocess(s) for s in bpe_strings]

        
        for i in range(0,len(real_topk_pos)):
            if real_topk_pos[i]>=1000:
                word_count[3]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    if bpe_strings[i+1].startswith('Ġ'):
                        para.add_run(' ')
                    elif bpe_strings[i+1].startswith('Ċ'):
                        para.add_run('\n')
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.VIOLET
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.VIOLET

            elif real_topk_pos[i]<1000 and real_topk_pos[i]>=100:
                word_count[2]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    if bpe_strings[i+1].startswith('Ġ'):
                        para.add_run(' ')
                    elif bpe_strings[i+1].startswith('Ċ'):
                        para.add_run('\n')
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.RED

            elif real_topk_pos[i]<100 and real_topk_pos[i]>=10:
                word_count[1]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    if bpe_strings[i+1].startswith('Ġ'):
                        para.add_run(' ')
                    elif bpe_strings[i+1].startswith('Ċ'):
                        para.add_run('\n')
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
            elif real_topk_pos[i]<10:
                word_count[0]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    if bpe_strings[i+1].startswith('Ġ'):
                        para.add_run(' ')
                    elif bpe_strings[i+1].startswith('Ċ'):
                        para.add_run('\n')
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        return word_count

    def gettext(self, file, fileName):
        text = ""
        if fileName.endswith('.docx'):
            text = docx2txt.process(file)
        elif fileName.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in range(len(reader.pages)):
                page_text = reader.pages[page]
                text+= page_text.extract_text()
        return text

    def split_text(self, project, text, para):
        word_count=[0,0,0,0]
        word_count_para=[0,0,0,0]
        max = len(text)
        i = 0
        while i < max:
            if i+1500 < max:
                word_count_para = project.lm.check_probabilities(text[i:i+1500], 40, para)
                i+=1500
            else:
                word_count_para = project.lm.check_probabilities(text[i:],40, para)
                i = max
            for j in range(len(word_count_para)):
                word_count[j] += word_count_para[j]
        return word_count
        

    def check_percentage(self,project,count):
        percent=['0','0','0','0']
        for i in range(0,4):
            percent[i] = str((count[i]/sum(count))*100)+'%'
        return percent

    def get_values(self, project, file, filename, topk, zip_files):
        doc = docx.Document()
        para = doc.add_paragraph('''''')

        text = project.lm.gettext(file, filename)
        file_name = [filename]
        if len(text)>2500:
            get_text = project.lm.split_text(project, text, para)
            count = project.lm.check_percentage(project,get_text)
        else:
            get_text = project.lm.check_probabilities(text,topk, para)
            count = project.lm.check_percentage(project,get_text)
        for j in range(len(count)):
                file_name.append(count[j])
        filename_docx = f"{filename}"
        if filename_docx.endswith('.pdf') or filename_docx.endswith('.txt'):
                filename_docx = filename_docx[:len(filename_docx)-4]+'.docx'
        doc.save(filename_docx)    
        zip_files.append(filename_docx)
        return file_name       

    def extract_files(self, project, file, topk=40):
        zip_files = []
        row=[['FileName','Top10','Top100','Top1000','Above1000']]
        if file.filename.endswith('.docx') or file.filename.endswith('.pdf'):
            output = project.lm.get_values(project, file, file.filename, topk, zip_files)
            row.append(output)
        elif zipfile.is_zipfile(file):
            with zipfile.ZipFile(file,'r') as zip:
                zip.extractall()
            for i in zip.infolist():
                if i.filename[0].isalpha():
                    if i.filename.endswith(".pdf") or i.filename.endswith(".docx") or i.filename.endswith(".txt") and 'MACOSX' not in i.filename:
                        output = project.lm.get_values(project, i.filename, i.filename, topk, zip_files)
                        row.append(output)
                    if i.filename.endswith(".docx") == False:
                        os.remove('./'+i.filename)
                    else:
                        print("No valid files in zip")
        else:
            print("Its not zip or pdf or docx or text file")
        
        output_file = "result.csv"
        with open(output_file, "w", newline="") as file:
            writer = csv.writer(file)
            writer.writerows(row)

        zip_files.append(output_file)
        in_memory_zip = io.BytesIO()
        with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
            for file in zip_files:
                zip_file.write(file)
                os.remove(file)

        in_memory_zip.seek(0)
        return send_file(in_memory_zip, download_name='result.zip', as_attachment=True)
   
    def postprocess(self, token):
        with_space = False
        with_break = False
        if token.startswith('Ġ'):
            with_space = True
            token = token[1:]
            # print(token)
        elif token.startswith('â'):
            token = ' '
        elif token.startswith('Ċ'):
            token = ' '
            with_break = True

        token = '-' if token.startswith('â') else token
        token = '“' if token.startswith('ľ') else token
        token = '”' if token.startswith('Ŀ') else token
        token = "'" if token.startswith('Ļ') else token

        if with_space:
            token = '\u0120' + token
        if with_break:
            token = '\u010A' + token

        return token

@register_api(name='BERT')
class BERTLM(AbstractLanguageChecker):
    def __init__(self, model_name_or_path="bert-base-cased"):
        super(BERTLM, self).__init__()
        self.device = torch.device(
            "cuda" if torch.cuda.is_available() else "cpu")
        self.tokenizer = BertTokenizer.from_pretrained(
            model_name_or_path,
            do_lower_case=False)
        self.model = BertForMaskedLM.from_pretrained(
            model_name_or_path)
        self.model.to(self.device)
        self.model.eval()
        # BERT-specific symbols
        self.mask_tok = self.tokenizer.convert_tokens_to_ids(["[MASK]"])[0]
        self.pad = self.tokenizer.convert_tokens_to_ids(["[PAD]"])[0]
        print("Loaded BERT model!")

    def check_probabilities(self, in_text, topk, para, max_context=20,
                            batch_size=20):
        '''
        Same behavior as GPT-2
        Extra param: max_context controls how many words should be
        fed in left and right
        Speeds up inference since BERT requires prediction word by word
        '''
        word_count =[0,0,0,0]
        in_text = "[CLS] " + in_text + " [SEP]"
        tokenized_text = self.tokenizer.tokenize(in_text)
        # Construct target
        y_toks = self.tokenizer.convert_tokens_to_ids(tokenized_text)
        # Only use sentence A embedding here since we have non-separable seq's
        segments_ids = [0] * len(y_toks)
        y = torch.tensor([y_toks]).to(self.device)
        segments_tensor = torch.tensor([segments_ids]).to(self.device)

        # TODO batching...
        # Create batches of (x,y)
        input_batches = []
        target_batches = []
        for min_ix in range(0, len(y_toks), batch_size):
            max_ix = min(min_ix + batch_size, len(y_toks) - 1)
            cur_input_batch = []
            cur_target_batch = []
            # Construct each batch
            for running_ix in range(max_ix - min_ix):
                tokens_tensor = y.clone()
                mask_index = min_ix + running_ix
                tokens_tensor[0, mask_index + 1] = self.mask_tok

                # Reduce computational complexity by subsetting
                min_index = max(0, mask_index - max_context)
                max_index = min(tokens_tensor.shape[1] - 1,
                                mask_index + max_context + 1)

                tokens_tensor = tokens_tensor[:, min_index:max_index]
                # Add padding
                needed_padding = max_context * 2 + 1 - tokens_tensor.shape[1]
                if min_index == 0 and max_index == y.shape[1] - 1:
                    # Only when input is shorter than max_context
                    left_needed = (max_context) - mask_index
                    right_needed = needed_padding - left_needed
                    p = torch.nn.ConstantPad1d((left_needed, right_needed),
                                               self.pad)
                    tokens_tensor = p(tokens_tensor)
                elif min_index == 0:
                    p = torch.nn.ConstantPad1d((needed_padding, 0), self.pad)
                    tokens_tensor = p(tokens_tensor)
                elif max_index == y.shape[1] - 1:
                    p = torch.nn.ConstantPad1d((0, needed_padding), self.pad)
                    tokens_tensor = p(tokens_tensor)

                cur_input_batch.append(tokens_tensor)
                cur_target_batch.append(y[:, mask_index + 1])
                # new_segments = segments_tensor[:, min_index:max_index]
            cur_input_batch = torch.cat(cur_input_batch, dim=0)
            cur_target_batch = torch.cat(cur_target_batch, dim=0)
            input_batches.append(cur_input_batch)
            target_batches.append(cur_target_batch)

        real_topk = []
        pred_topk = []

        with torch.no_grad():
            for src, tgt in zip(input_batches, target_batches):
                # Compute one batch of inputs
                # By construction, MASK is always the middle
                logits = self.model(src, torch.zeros_like(src))[0][:,
                         max_context + 1]
                yhat = torch.softmax(logits, dim=-1)

                sorted_preds = np.argsort(-yhat.data.cpu().numpy())
                # TODO: compare with batch of tgt

                # [(pos, prob), ...]
                real_topk_pos = list(
                    [int(np.where(sorted_preds[i] == tgt[i].item())[0][0])
                     for i in range(yhat.shape[0])])
                real_topk_probs = yhat[np.arange(
                    0, yhat.shape[0], 1), tgt].data.cpu().numpy().tolist()
                real_topk.extend(list(zip(real_topk_pos, real_topk_probs)))

                # # [[(pos, prob), ...], [(pos, prob), ..], ...]
                pred_topk.extend([list(zip(self.tokenizer.convert_ids_to_tokens(
                    sorted_preds[i][:topk]),
                    yhat[i][sorted_preds[i][
                            :topk]].data.cpu().numpy().tolist()))
                    for i in range(yhat.shape[0])])

        bpe_strings = [self.postprocess(s) for s in tokenized_text]
        pred_topk = [[(self.postprocess(t[0]), t[1]) for t in pred] for pred in pred_topk]

        for i in range(0,len(real_topk)-1):
            if bpe_strings[i+1].startswith('Ġ'):
                para.add_run(' ')
            elif bpe_strings[i+1].startswith('Ċ'):
                para.add_run('\n')
            if real_topk[i][0]>=1000:
                word_count[3]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.VIOLET
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.VIOLET
            elif real_topk[i][0]<1000 and real_topk[i][0]>=100:
                word_count[2]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.RED
            elif real_topk[i][0]<100 and real_topk[i][0]>=10:
                word_count[1]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif real_topk[i][0]<10:
                word_count[0]+=1
                if 'Ġ' in bpe_strings[i+1] or 'Ċ' in bpe_strings[i+1]:
                    para.add_run(bpe_strings[i+1][1:]).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    para.add_run(bpe_strings[i+1]).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        return word_count
    
    def gettext(self, file, fileName):
        text = ""
        if fileName.endswith('.docx'):
            text = docx2txt.process(file)
        elif fileName.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in range(len(reader.pages)):
                page_text = reader.pages[page]
                text+= page_text.extract_text()
        text = re.sub(r'[^a-zA-Z\d\s]', '', text)
        text = text.strip()
        return text

    def check_percentage(self,project,count):
        percent=['0','0','0','0']
        for i in range(0,4):
            percent[i] = str((count[i]/sum(count))*100)+'%'
        return percent

    def get_values(self, project, file, filename, topk, zip_files):
        doc = docx.Document()
        para = doc.add_paragraph('''''')

        text = project.lm.gettext(file, filename)
        file_name = [filename]
        get_text = project.lm.check_probabilities(text,topk,para)
        count = project.lm.check_percentage(project,get_text)
        for j in range(len(count)):
                file_name.append(count[j])
        filename_docx = f"{filename}"
        if filename_docx.endswith('.pdf') or filename_docx.endswith('.txt'):
                filename_docx = filename_docx[:len(filename_docx)-4]+'.docx'
        doc.save(filename_docx)    
        zip_files.append(filename_docx)
        return file_name       

    def extract_files(self, project, file, topk=20):
        zip_files = []
        row=[['FileName','Top10','Top100','Top1000','Above1000']]
        if file.filename.endswith('.docx') or file.filename.endswith('.pdf'):
            output = project.lm.get_values(project, file, file.filename, topk, zip_files)
            row.append(output)
        elif zipfile.is_zipfile(file):
            with zipfile.ZipFile(file,'r') as zip:
                zip.extractall()
            for i in zip.infolist():
                if i.filename[0].isalpha() == True:
                    if i.filename.endswith(".pdf") or i.filename.endswith(".docx") or i.filename.endswith(".txt") and 'MACOSX' not in i.filename:
                        output = project.lm.get_values(project, i.filename, i.filename, topk, zip_files)
                        row.append(output)
                    if i.filename.endswith(".docx") == False:
                        os.remove('./'+i.filename)
                    else:
                        print("No valid files in zip")
        else:
            print("Its not zip or pdf or docx or text file")
        
        output_file = "result.csv"
        with open(output_file, "w", newline="") as file:
            writer = csv.writer(file)
            writer.writerows(row)

        zip_files.append(output_file)
        in_memory_zip = io.BytesIO()
        with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
            for file in zip_files:
                zip_file.write(file)
                os.remove(file)

        in_memory_zip.seek(0)
        return send_file(in_memory_zip, download_name='result.zip', as_attachment=True)

    def postprocess(self, token):

        with_space = True
        with_break = token == '[SEP]'
        if token.startswith('##'):
            with_space = False
            token = token[2:]

        if with_space:
            token = '\u0120' + token
        if with_break:
            token = '\u010A' + token
        #
        # # print ('....', token)
        return token


def main():
    raw_text = """"""
    