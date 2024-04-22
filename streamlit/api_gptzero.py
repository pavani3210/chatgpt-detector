import csv
import io
import json
import os
import zipfile
import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX
import requests


def extract_files(file):
    zip_files = []
    row=[['FileName','Status']]
    if file.name.endswith('.docx') or file.name.endswith('.pdf') or file.name.endswith('.txt'):
        output_gpt = get_values(file.name, file.read(), file.type, zip_files)
        row.append(output_gpt)
    elif zipfile.is_zipfile(file):
        count_pdf_docx = 0
        with zipfile.ZipFile(file, 'r') as zip_file:
            for file_name in zip_file.namelist():
                if file_name[0].isalpha() and (file_name.endswith(".pdf") or file_name.endswith(".docx") or file_name.endswith(".txt")) and 'MACOSX' not in file_name:
                    with zip_file.open(file_name) as file:
                        count_pdf_docx += 1
                        if file_name.endswith(".docx"):
                            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        if file_name.endswith('.pdf'):
                            content_type = "application/pdf"
                        file_stream = file.read()
                        output_gpt = get_values(file_name, file_stream, content_type, zip_files)
                        row.append(output_gpt)
        if count_pdf_docx == 0:
            st.write("No valid files in zip")
    else:
        st.write("Its not zip or pdf or docx or text file")
    
    output_file = "result.csv"
    with open(output_file, "w", newline="") as file:
        writer = csv.writer(file)
        writer.writerows(row)

    zip_files.append(output_file)
    in_memory_zip = io.BytesIO()
    with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
        for file in zip_files:
            zip_file.write(file)
            os.remove(file)

    in_memory_zip.seek(0)
    return in_memory_zip.getvalue()

def get_values(filename, stream, content_type, zip_files):
    doc = docx.Document()
    para = doc.add_paragraph('''''')
    file_name = [filename]
    filename_docx = f"{filename}"
    if filename.endswith('.docx'):
        filename_docx = filename[:len(filename_docx)-5]+'.docx'
    else:
        filename_docx = filename[:len(filename_docx)-4]+'.docx'
    status = check_gptzero(filename, stream, content_type, para) 
    if status == 0:
        file_name.append('Human Written')
    elif status < 0.5:
        file_name.append('Most likely Human Written')
    elif status < 1:
        file_name.append('Most likely AI/GPT Generated')
    elif status == 1:
        file_name.append('AI/GPT Generated')
 
    doc.save(filename_docx)    
    zip_files.append(filename_docx)
    return file_name       

def check_gptzero(filename,stream,content_type, para):
    API_URL = 'https://api.gptzero.me/v2/predict/files'
    api_key = 'bedffe57940747328050de84c2daddd3'
    headers = {
    'x-api-key': api_key,
    'Accept': 'application/json'
    }
    form_data = {
        'files': (filename, stream, content_type)
    }  
    try:
        response = requests.post(API_URL, headers=headers, files = form_data)
        decoded_content = response.content.decode('utf-8')
        json_content = json.loads(decoded_content)
        for i in json_content['documents']:
            status = i['average_generated_prob']
            for j in range(len(i['sentences'])):
                if i['sentences'][j]['generated_prob'] > 0.5:
                    para.add_run(i['sentences'][j]['sentence']+'\n').font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    para.add_run(i['sentences'][j]['sentence']+'\n')
    except requests.exceptions.RequestException as e:
        st.error(f"Error from GPT Zero API")
        status = -1  
    except json.JSONDecodeError as e:
        st.error(f"Error decoded Please use other valid document")
        status = -1  
    return status
    